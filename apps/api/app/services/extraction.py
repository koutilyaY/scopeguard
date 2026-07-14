"""Document text extraction (PDF via PyMuPDF, DOCX via python-docx, TXT/EML).

Extraction preserves page numbers and heading/section hints where the format allows.
Documents are treated as untrusted data: no macros or embedded content are executed.
"""

import email
import email.policy
import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger("scopeguard.extraction")

UNREADABLE_MESSAGE = (
    "This document does not contain sufficient machine-readable text. "
    "OCR is not enabled in this version."
)

# A PDF page yielding fewer characters than this is considered image-only.
MIN_CHARS_PER_READABLE_DOC = 40


@dataclass
class ExtractedPage:
    page_number: int
    text: str


@dataclass
class ExtractionResult:
    ok: bool
    pages: list[ExtractedPage] = field(default_factory=list)
    error: str | None = None

    @property
    def full_text(self) -> str:
        return "\n\n".join(f"[page {p.page_number}]\n{p.text}" for p in self.pages)

    @property
    def plain_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)


def extract_pdf(data: bytes) -> ExtractionResult:
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        return ExtractionResult(ok=False, error=f"Could not open PDF: {exc}")
    pages: list[ExtractedPage] = []
    try:
        for index, page in enumerate(doc):
            text = page.get_text("text").strip()
            pages.append(ExtractedPage(page_number=index + 1, text=text))
    finally:
        doc.close()
    total_chars = sum(len(p.text) for p in pages)
    if total_chars < MIN_CHARS_PER_READABLE_DOC:
        return ExtractionResult(ok=False, pages=pages, error=UNREADABLE_MESSAGE)
    return ExtractionResult(ok=True, pages=pages)


def extract_docx(data: bytes) -> ExtractionResult:
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        return ExtractionResult(ok=False, error=f"Could not open DOCX: {exc}")
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            lines.append(f"\n## {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    text = "\n".join(lines)
    if len(text.strip()) < MIN_CHARS_PER_READABLE_DOC:
        return ExtractionResult(ok=False, error=UNREADABLE_MESSAGE)
    # DOCX has no fixed pagination; expose as a single logical page.
    return ExtractionResult(ok=True, pages=[ExtractedPage(page_number=1, text=text)])


def extract_txt(data: bytes) -> ExtractionResult:
    try:
        text = data.decode("utf-8", errors="replace").strip()
    except Exception as exc:  # pragma: no cover - decode with replace shouldn't raise
        return ExtractionResult(ok=False, error=str(exc))
    if len(text) < 1:
        return ExtractionResult(ok=False, error=UNREADABLE_MESSAGE)
    return ExtractionResult(ok=True, pages=[ExtractedPage(page_number=1, text=text)])


@dataclass
class ParsedEmail:
    subject: str
    sender: str | None
    recipients: str | None
    date: str | None
    body: str


def extract_eml(data: bytes) -> tuple[ExtractionResult, ParsedEmail | None]:
    try:
        message = email.message_from_bytes(data, policy=email.policy.default)
    except Exception as exc:
        return ExtractionResult(ok=False, error=f"Could not parse email: {exc}"), None
    body = ""
    text_part = message.get_body(preferencelist=("plain", "html"))
    if text_part is not None:
        content = text_part.get_content()
        if text_part.get_content_type() == "text/html":
            from bs4 import BeautifulSoup

            body = BeautifulSoup(content, "html.parser").get_text("\n")
        else:
            body = content
    parsed = ParsedEmail(
        subject=str(message.get("Subject", "")).strip() or "(no subject)",
        sender=str(message.get("From", "")) or None,
        recipients=str(message.get("To", "")) or None,
        date=str(message.get("Date", "")) or None,
        body=body.strip(),
    )
    text = f"Subject: {parsed.subject}\nFrom: {parsed.sender}\nTo: {parsed.recipients}\n\n{parsed.body}"
    return ExtractionResult(ok=True, pages=[ExtractedPage(page_number=1, text=text)]), parsed


def extract_document(extension: str, data: bytes) -> ExtractionResult:
    if extension == ".pdf":
        return extract_pdf(data)
    if extension == ".docx":
        return extract_docx(data)
    if extension in (".txt", ".csv"):
        return extract_txt(data)
    if extension == ".eml":
        result, _ = extract_eml(data)
        return result
    return ExtractionResult(ok=False, error=f"No extractor for {extension}")
